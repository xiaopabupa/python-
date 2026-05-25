from docx import Document
import sys

doc = Document(sys.argv[1])

# Check all body elements
print("=== BODY XML summary ===")
body = doc.element.body
for child in body:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    text = ''.join(child.itertext()).strip()
    if text:
        print(f"  {tag}: {text[:200]}")

# Check sections
print("\n=== SECTIONS ===")
for i, section in enumerate(doc.sections):
    print(f"Section {i}:")
    header = section.header
    footer = section.footer
    for p in header.paragraphs:
        if p.text.strip():
            print(f"  Header: {p.text[:200]}")
    for p in footer.paragraphs:
        if p.text.strip():
            print(f"  Footer: {p.text[:200]}")

# Check if there are inline shapes or other elements
print("\n=== ALL PARAGRAPHS (including empty) ===")
for i, p in enumerate(doc.paragraphs):
    print(f"[P{i}] style={p.style.name} len={len(p.text)} text='{p.text[:200]}'")

# Check for tables count
print(f"\n=== TABLES COUNT: {len(doc.tables)} ===")
for i, t in enumerate(doc.tables):
    print(f"Table {i}: {len(t.rows)} rows x {len(t.columns)} cols")
